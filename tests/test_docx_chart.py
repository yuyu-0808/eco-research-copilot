from docx import Document
from docx.shared import Inches
from docx.chart.data import CategoryChartData
from docx.enum.chart import XL_CHART_TYPE
import os

def test_native_word_chart():
    print("⏳ 正在生成原生可编辑 Word 报告...")
    
    # 1. 初始化 Word 文档
    doc = Document()
    doc.add_heading('Eco-Research 宏观商业分析报告 (Demo)', level=1)
    doc.add_paragraph('✅ 核心亮点：以下图表由 AI 根据宏观经济数据自动生成，【双击图表】即可唤起 Excel 编辑底层数据。')

    # 2. 模拟大模型（Agent）清洗出来的结构化 JSON 数据
    chart_data = CategoryChartData()
    chart_data.categories = ['比亚迪 (BYD)', '特斯拉 (Tesla)', '长城 (GWM)', '吉利 (Geely)']
    # 添加数据序列（比如销量预测）
    chart_data.add_series('2026年马来西亚销量预测 (万辆)', (4.2, 2.1, 1.5, 1.8))

    # 3. 在 Word 中插入原生的柱状图
    x, y, cx, cy = Inches(1), Inches(1), Inches(6), Inches(4)
    doc.add_chart(
        XL_CHART_TYPE.COLUMN_CLUSTERED, x, y, cx, cy, chart_data
    )
    
    # 4. 获取当前项目根目录，将报告保存在最外面方便查看# (因为脚本在 tests 文件夹里，我们让它往上走一级存到根目录)
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(current_dir)
    output_filename = os.path.join(project_root, 'Eco_Research_Demo.docx')
    
    doc.save(output_filename)
    print(f"🎉 成功！请在左侧目录查看生成的：Eco_Research_Demo.docx")

if __name__ == "__main__":
    test_native_word_chart()