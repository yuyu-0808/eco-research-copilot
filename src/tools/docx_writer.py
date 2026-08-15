import os
import re
import json
import urllib.parse
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE

from src.utils.config import Config


class DocxWriter:
    # 学术论文格式常量
    CN_FONT = '宋体'
    EN_FONT = 'Times New Roman'
    HEADING_FONT = '黑体'

    def __init__(self, project_name: str):
        root_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        project_dir = os.path.join(root_dir, "projects", project_name)
        os.makedirs(project_dir, exist_ok=True)
        self.project_dir = project_dir
        self.output_docx = os.path.join(project_dir, '05_final_report.docx')

    # ---------- 格式工具 ----------
    @staticmethod
    def _set_run_font(run, size=12, bold=False, cn='宋体', en='Times New Roman', color=None):
        run.font.name = en
        run.font.size = Pt(size)
        run.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), en)
        rfonts.set(qn('w:hAnsi'), en)
        rfonts.set(qn('w:eastAsia'), cn)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    @staticmethod
    def _set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=2, space_after=0):
        pf = p.paragraph_format
        pf.alignment = align
        pf.line_spacing = line_spacing
        pf.space_after = Pt(space_after)
        if first_indent_chars:
            ind = p._p.get_or_add_pPr().get_or_add_ind()
            ind.set(qn('w:firstLineChars'), str(first_indent_chars * 100))

    # ---------- 元素添加 ----------
    def _add_title(self, doc, text):
        """报告主标题：黑体、居中、加粗、较大字号"""
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=12)
        run = p.add_run(text)
        self._set_run_font(run, size=16, bold=True, cn=self.HEADING_FONT)

    def _add_meta(self, doc, text):
        """元信息（发布日期等）：居中、小字"""
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=10)
        run = p.add_run(text)
        self._set_run_font(run, size=10.5, cn=self.CN_FONT)

    def _add_h1(self, doc, text):
        """一级标题：黑体 15pt、居中、加粗"""
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=6)
        run = p.add_run(text)
        self._set_run_font(run, size=15, bold=True, cn=self.HEADING_FONT)

    def _add_h2(self, doc, text):
        """二级标题：宋体 14pt、左对齐、加粗"""
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_chars=2, space_after=3)
        run = p.add_run(text)
        self._set_run_font(run, size=14, bold=True, cn=self.CN_FONT)

    def _add_body(self, doc, text):
        """正文段落：宋体 12pt、1.5倍行距、两端对齐、首行缩进2字符，支持 **加粗**"""
        p = doc.add_paragraph()
        self._set_para(p)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = p.add_run(part[2:-2])
                self._set_run_font(run, size=12, bold=True)
            else:
                run = p.add_run(part)
                self._set_run_font(run, size=12)

    def _add_caption(self, doc, text):
        """图表题注：居中、加粗、小字"""
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=4)
        run = p.add_run(text)
        self._set_run_font(run, size=10.5, bold=True, cn=self.CN_FONT)

    def _add_hyperlink(self, paragraph, url, text):
        """在段落中插入可点击超链接"""
        if not url:
            run = paragraph.add_run(text)
            self._set_run_font(run, size=10.5)
            return
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rpr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), self.CN_FONT)
        rpr.append(rfonts)
        new_run.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # ---------- 图表生成 ----------
    def _build_chart_config(self, chart_type, labels, data, title, legend_label=None):
        colors = ["#5470C6", "#91CC75", "#FAC858", "#EE6666", "#73C0DE", "#3BA272", "#FC8452", "#9A60B4"]
        legend = legend_label or title  # 图例用 label（指标+单位），为空则回退 title
        dataset = {"data": data}
        if chart_type == "line":
            dataset.update({"label": legend, "borderColor": "#5470C6",
                            "backgroundColor": "rgba(84,112,198,0.25)", "fill": False, "tension": 0.3})
        elif chart_type == "bar":
            dataset.update({"label": legend, "backgroundColor": colors[:max(1, len(data))],
                            "borderColor": "#5470C6", "borderWidth": 1})
        elif chart_type == "pie":
            dataset.update({"backgroundColor": colors[:max(1, len(data))]})
        config = {
            "type": chart_type,
            "data": {"labels": labels, "datasets": [dataset]},
            "options": {
                "plugins": {"legend": {"display": True, "position": "bottom"}},
                "scales": {"y": {"beginAtZero": True}} if chart_type != "pie" else {}
            }
        }
        return config

    def _add_chart(self, doc, chart, index):
        """生成一张图（QuickChart）并插入，返回临时图片路径"""
        chart_type = chart.get("type", "line")
        if chart_type not in ("line", "bar", "pie"):
            chart_type = "line"
        title = chart.get("title", f"图{index}")
        legend_label = chart.get("label", "")
        labels = chart.get("labels", [])
        data = chart.get("data", [])
        if not data:
            return None

        img_path = os.path.join(self.project_dir, f'temp_chart_{index}.png')
        try:
            config = self._build_chart_config(chart_type, labels, data, title, legend_label)
            encoded = urllib.parse.quote(json.dumps(config, ensure_ascii=False))
            resp = requests.get(f"https://quickchart.io/chart?c={encoded}", timeout=30)
            with open(img_path, 'wb') as f:
                f.write(resp.content)
            if not resp.content:
                return None
        except Exception:
            return None

        self._add_caption(doc, title)
        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=12)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        return img_path

    def _add_table(self, doc, table_data, index):
        """插入一张表（Word 原生表格）"""
        title = table_data.get("title", f"表{index}")
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if not headers and not rows:
            return
        self._add_caption(doc, title)

        if headers:
            n_cols = len(headers)
            tbl = doc.add_table(rows=1, cols=n_cols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            for j, h in enumerate(headers):
                hdr[j].text = str(h)
                for para in hdr[j].paragraphs:
                    for run in para.runs:
                        self._set_run_font(run, size=10.5, bold=True)
            for row in rows:
                cells = tbl.add_row().cells
                for j in range(n_cols):
                    cells[j].text = str(row[j]) if j < len(row) else ""
                    for para in cells[j].paragraphs:
                        for run in para.runs:
                            self._set_run_font(run, size=10.5)

        p = doc.add_paragraph()
        self._set_para(p, first_indent_chars=0, space_after=6)

    # ---------- 正文渲染 ----------
    def _add_markdown(self, doc, markdown_text, charts, tables, temp_files):
        """把 Markdown 正文渲染成论文格式段落，解析 [[CHART:n]]/[[TABLE:n]] 占位符就地插图"""
        used_charts = set()
        used_tables = set()
        if markdown_text:
            for line in markdown_text.strip().split('\n'):
                s = line.strip()
                if not s:
                    continue
                # 图表占位符：就地插入
                m = re.match(r'\[\[CHART:(\d+)\]\]', s)
                if m:
                    idx = int(m.group(1)) - 1
                    if charts and 0 <= idx < len(charts):
                        f = self._add_chart(doc, charts[idx], idx + 1)
                        if f:
                            temp_files.append(f)
                        used_charts.add(idx)
                    continue
                m = re.match(r'\[\[TABLE:(\d+)\]\]', s)
                if m:
                    idx = int(m.group(1)) - 1
                    if tables and 0 <= idx < len(tables):
                        self._add_table(doc, tables[idx], idx + 1)
                        used_tables.add(idx)
                    continue
                # 标题 / 正文
                if s.startswith('### '):
                    self._add_h2(doc, s[4:].strip())
                elif s.startswith('## '):
                    self._add_h1(doc, s[3:].strip())
                elif s.startswith('# '):
                    self._add_h1(doc, s[2:].strip())
                elif s.startswith('- ') or s.startswith('* '):
                    self._add_body(doc, s[2:].strip())
                elif re.match(r'^\d+[.)]\s', s):
                    self._add_body(doc, re.sub(r'^\d+[.)]\s*', '', s))
                else:
                    self._add_body(doc, s)
        return used_charts, used_tables

    def _add_references(self, doc, references):
        """渲染参考文献列表"""
        refs = references or []
        if not refs:
            return
        self._add_h1(doc, '参考文献')
        for ref in refs:
            idx = ref.get('index', '')
            title = ref.get('title', '') or ''
            url = ref.get('url', '') or ''
            p = doc.add_paragraph()
            self._set_para(p, first_indent_chars=0)
            run = p.add_run(f'[{idx}] ')
            self._set_run_font(run, size=10.5, bold=True)
            self._add_hyperlink(p, url, title)

    # ---------- 券商研报排版（封面 / 目录 / 页眉页脚 / 数据附表） ----------
    def _setup_header_footer(self, doc, title: str):
        """页眉=自定义页眉或报告标题，页脚=自定义页脚或居中页码。"""
        section = doc.sections[0]
        header_text = Config.REPORT_HEADER or title
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(header_text)
        self._set_run_font(run, size=9, cn=self.CN_FONT, color="8B92AE")
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if Config.REPORT_FOOTER:
            fr = fp.add_run(Config.REPORT_FOOTER)
            self._set_run_font(fr, size=9, cn=self.CN_FONT, color="8B92AE")
        else:
            self._add_page_number(fp)

    def _add_page_number(self, paragraph):
        """在段落里插入 PAGE 字段（自动页码）。"""
        run = paragraph.add_run()
        self._set_run_font(run, size=9, cn=self.CN_FONT, color="8B92AE")
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    def _add_cover(self, doc, ai_data: dict):
        """封面页：Logo（可选）+ 报告标题 + 副标题 + 日期 + 免责声明。"""
        title = ai_data.get("report_title", "行业研究报告")
        for _ in range(2):
            p = doc.add_paragraph()
            self._set_para(p, first_indent_chars=0)

        # 自定义 Logo（若配置且文件存在）
        logo = (Config.REPORT_LOGO or "").strip()
        if logo and os.path.exists(logo):
            try:
                p = doc.add_paragraph()
                self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=14)
                run = p.add_run()
                run.add_picture(logo, width=Inches(1.6))
            except Exception:
                pass

        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=18)
        run = p.add_run(title)
        self._set_run_font(run, size=22, bold=True, cn=self.HEADING_FONT)

        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=10)
        run = p.add_run("行业研究报告")
        self._set_run_font(run, size=15, cn=self.CN_FONT, color="6A7180")

        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=8)
        run = p.add_run(f"发布日期：{ai_data.get('publish_date', '')}    研究引擎：Eco-Research Copilot")
        self._set_run_font(run, size=11, cn=self.CN_FONT, color="6A7180")

        for _ in range(6):
            doc.add_paragraph()

        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_chars=0, space_after=4)
        run = p.add_run("免责声明")
        self._set_run_font(run, size=10.5, bold=True, cn=self.CN_FONT)

        p = doc.add_paragraph()
        self._set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_chars=0, space_after=0)
        disclaimer = Config.REPORT_DISCLAIMER or (
            "本报告由 Eco-Research Copilot 自动生成，仅供研究参考，不构成任何投资建议。"
            "报告数据来自公开信源，虽经多级代码校验，仍可能存在遗漏或偏差，使用者应自行核实。"
        )
        run = p.add_run(disclaimer)
        self._set_run_font(run, size=9, cn=self.CN_FONT, color="8B92AE")
        doc.add_page_break()

    def _add_toc(self, doc, ai_data: dict):
        """静态目录：基于正文二级标题。"""
        headings = self._extract_headings(ai_data.get("markdown_report", ""))
        if not headings:
            return
        self._add_h1(doc, "目录")
        for h in headings:
            p = doc.add_paragraph()
            self._set_para(p, first_indent_chars=0, space_after=3)
            run = p.add_run(h)
            self._set_run_font(run, size=11, cn=self.CN_FONT)
        doc.add_page_break()

    @staticmethod
    def _extract_headings(md: str) -> list:
        if not md:
            return []
        return re.findall(r'^##\s+(.+)$', md, re.M)

    def _add_data_appendix(self, doc, evidence):
        """数据附表：汇总报告中所有量化指标成规范表格，方便直接取用。"""
        from src.utils.metrics_store import extract_metrics
        entries = extract_metrics(evidence or [])
        if not entries:
            return
        self._add_h1(doc, "附：数据附表")
        headers = ["指标", "数值", "单位", "时间", "信源等级", "来源"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for j, h in enumerate(headers):
            hdr[j].text = h
            for para in hdr[j].paragraphs:
                for run in para.runs:
                    self._set_run_font(run, size=10.5, bold=True)
        for en in entries:
            cells = tbl.add_row().cells
            row = [
                en.get("metric_label", ""),
                str(en.get("value", "")),
                en.get("unit", "") or "",
                en.get("period", "") or "",
                en.get("source_tier", ""),
                en.get("source_title") or en.get("publisher", "") or "",
            ]
            for j, v in enumerate(row):
                cells[j].text = v
                for para in cells[j].paragraphs:
                    for run in para.runs:
                        self._set_run_font(run, size=10.5)
        p = doc.add_paragraph()
        self._set_para(p, first_indent_chars=0, space_after=6)

    # ---------- 主流程 ----------
    def generate_report(self, ai_data: dict) -> str:
        """从头构建一份券商研报格式的研究报告（封面 / 目录 / 正文 / 数据附表 / 参考文献）。"""
        doc = Document()
        temp_files = []
        title = ai_data.get("report_title", "行业研究报告")

        # 0. 页眉页脚（券商研报排版）
        self._setup_header_footer(doc, title)

        # 1. 封面 + 免责声明
        self._add_cover(doc, ai_data)

        # 2. 目录
        self._add_toc(doc, ai_data)

        # 3. 摘要（核心洞察）
        core = ai_data.get("core_insights", "")
        if core:
            self._add_h1(doc, '摘要')
            self._add_body(doc, core)

        # 4. 正文（图表穿插其中）
        tables = ai_data.get("tables", [])
        charts = ai_data.get("charts", [])
        md = ai_data.get("markdown_report", "")
        used_charts, used_tables = self._add_markdown(doc, md, charts, tables, temp_files)

        # 4.5 兜底：正文未引用的图表/表格追加到「附：补充数据」
        unused_charts = [c for i, c in enumerate(charts) if i not in used_charts]
        unused_tables = [t for i, t in enumerate(tables) if i not in used_tables]
        if unused_tables or unused_charts:
            self._add_h1(doc, '附：补充数据')
            for i, t in enumerate(tables):
                if i not in used_tables:
                    self._add_table(doc, t, i + 1)
            for i, c in enumerate(charts):
                if i not in used_charts:
                    f = self._add_chart(doc, c, i + 1)
                    if f:
                        temp_files.append(f)

        # 5. 数据附表（量化指标汇总）
        self._add_data_appendix(doc, ai_data.get("evidence", []))

        # 6. 参考文献
        self._add_references(doc, ai_data.get("references", []))

        doc.save(self.output_docx)

        # 清理临时图片
        for f in temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except Exception:
                pass

        return self.output_docx
